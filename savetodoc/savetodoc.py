from docx import Document
from docx.shared import Inches

def savetodoc(promptcontainer, docname):
    doc = Document()
    doc.add_heading(docname)
    for prompt in promptcontainer.customPromptList:
        doc.add_heading(prompt.prompttext)
        doc.add_paragraph(prompt.promptresponse)
    doc.save('demo.docx')
        
